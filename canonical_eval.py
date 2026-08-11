"""Deterministic end-to-end reliability benchmark for the canonical QA path."""

from __future__ import annotations

import io
from collections import defaultdict
from dataclasses import asdict, dataclass
from typing import Any

import fitz
from docx import Document
from openpyxl import Workbook

from canonical_rag import ActiveDocumentService, answer_question
from document_normalizer import normalize_document


@dataclass(frozen=True)
class EvaluationCase:
    format: str
    question: str
    expected: str
    expected_status: str = "ANSWER"


def _facts() -> list[list[str]]:
    return [
        ["System", "Host", "Protocol", "Directory", "Frequency", "Description", "Transformation", "Password",
         "Version", "Status", "Duplicate mechanism", "Archive directory"],
        ["BI", "10.0.0.1", "SFTP", "/data/bi", "5 minutes",
         "The system transfers files directly using PUSH.", "No transformation is applied.", "real-secret",
         "2.1", "Processed", "Checksum", "/archive/bi"],
    ]


def _documents() -> dict[str, bytes]:
    rows = _facts()
    csv_bytes = (",".join(rows[0]) + "\n" + ",".join(rows[1]) + "\n").encode()

    workbook = Workbook()
    for row in rows:
        workbook.active.append(row)
    xlsx = io.BytesIO()
    workbook.save(xlsx)

    word = Document()
    word.add_heading("Operations", level=1)
    table = word.add_table(rows=2, cols=len(rows[0]))
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    docx = io.BytesIO()
    word.save(docx)

    pdf = fitz.open()
    page = pdf.new_page()
    text = " | ".join(f"{label} = {value}" for label, value in zip(*rows))
    page.insert_textbox(fitz.Rect(36, 36, 560, 780), text, fontsize=9)
    pdf_bytes = pdf.tobytes()
    pdf.close()
    return {"docx": docx.getvalue(), "pdf": pdf_bytes, "xlsx": xlsx.getvalue(), "csv": csv_bytes}


def _grounded_llm(prompt: str) -> dict[str, Any]:
    import re

    ids = re.findall(r"\[([a-f0-9]{64})\]", prompt)
    query = prompt.split("QUESTION:", 1)[-1].split("\nEVIDENCE:", 1)[0].casefold()
    if "send" in query or "transfer" in query or "envoie" in query:
        return {"answer": "PUSH", "supporting_evidence_ids": ids[:1], "supported": bool(ids)}
    if "transform" in query:
        return {"answer": "No", "supporting_evidence_ids": ids[:1], "supported": bool(ids)}
    return {"answer": "", "supporting_evidence_ids": [], "supported": False}


def run_benchmark() -> dict[str, Any]:
    cases: list[EvaluationCase] = []
    for file_format in ("docx", "pdf", "xlsx", "csv"):
        cases.extend([
            EvaluationCase(file_format, "What is the BI host?", "10.0.0.1"),
            EvaluationCase(file_format, "What is the BI protocol?", "SFTP"),
            EvaluationCase(file_format, "What is the BI directory?", "/data/bi"),
            EvaluationCase(file_format, "What is the collection frequency?", "5 minutes"),
            EvaluationCase(file_format, "How does the system send files?", "PUSH"),
            EvaluationCase(file_format, "Are files transformed?", "No"),
            EvaluationCase(file_format, "What is the version?", "2.1"),
            EvaluationCase(file_format, "What is the processing status?", "Processed"),
            EvaluationCase(file_format, "What is the duplicate mechanism?", "Checksum"),
            EvaluationCase(file_format, "What is the archive directory?", "/archive/bi"),
            EvaluationCase(file_format, "Quel est le protocole de BI ?", "SFTP"),
            EvaluationCase(file_format, "What is the BI database?", "NO_EXPLICIT_EVIDENCE", "NO_EVIDENCE"),
            EvaluationCase(file_format, "What is the password?", "NO_EXPLICIT_EVIDENCE", "SENSITIVE_BLOCK"),
        ])
    metrics: dict[str, Any] = {
        "total_cases": len(cases), "factual_answerable_cases": 0, "correct": 0, "wrong": 0,
        "false_no_evidence": 0, "ambiguous_refusal": 0, "sensitive_block": 0,
        "cross_document_leakage": 0, "secret_leakage": 0, "unsupported_generated_answer": 0,
        "deterministic_contribution": 0, "grounded_fallback_contribution": 0,
    }
    per_format: dict[str, dict[str, int]] = defaultdict(lambda: {"total": 0, "correct": 0})
    documents = _documents()
    for case in cases:
        document = normalize_document(documents[case.format], f"benchmark.{case.format}")
        context = ActiveDocumentService().select(document)
        result = answer_question(context, case.question)
        metrics["factual_answerable_cases"] += case.expected_status == "ANSWER"
        per_format[case.format]["total"] += 1
        correct = result.status == case.expected_status and result.answer == case.expected
        if correct:
            metrics["correct"] += 1
            per_format[case.format]["correct"] += 1
        else:
            metrics["wrong"] += result.status == "ANSWER"
            metrics["false_no_evidence"] += case.expected_status == "ANSWER" and result.status != "ANSWER"
        metrics["sensitive_block"] += result.status == "SENSITIVE_BLOCK"
        metrics["ambiguous_refusal"] += result.status == "AMBIGUOUS"
        metrics["cross_document_leakage"] += sum(block.file_hash != context.file_hash for block in result.evidence_blocks)
        metrics["secret_leakage"] += "real-secret" in result.answer
        metrics["deterministic_contribution"] += result.method == "deterministic_structured"
        metrics["grounded_fallback_contribution"] += 0
        metrics.setdefault("local_span_contribution", 0)
        metrics["local_span_contribution"] += result.method == "local_span_extraction"
    # Adversarial ambiguity and document-switch cases exercise the same public
    # end-to-end path but are not attributed to one source format metric.
    conflict = normalize_document(b"System,Host\nBI,1.1.1.1\nBI,2.2.2.2\n", "conflict.csv")
    conflict_result = answer_question(ActiveDocumentService().select(conflict), "What is the BI host?")
    metrics["total_cases"] += 1
    metrics["correct"] += conflict_result.status == "AMBIGUOUS"
    metrics["ambiguous_refusal"] += conflict_result.status == "AMBIGUOUS"

    switch_service = ActiveDocumentService()
    switch_service.select(normalize_document(b"System,Host\nBI,1.1.1.1\n", "a.csv"))
    active_b = switch_service.select(normalize_document(b"System,Host\nBI,2.2.2.2\n", "b.csv"))
    switch_result = answer_question(active_b, "What is the BI host?")
    metrics["total_cases"] += 1
    metrics["factual_answerable_cases"] += 1
    switch_correct = switch_result.answer == "2.2.2.2" and all(
        block.file_hash == active_b.file_hash for block in switch_result.evidence_blocks
    )
    metrics["correct"] += switch_correct
    metrics["wrong"] += not switch_correct
    metrics["cross_document_leakage"] += sum(block.file_hash != active_b.file_hash for block in switch_result.evidence_blocks)

    metrics["per_format"] = {
        key: {**value, "accuracy": value["correct"] / value["total"] if value["total"] else 0.0}
        for key, value in per_format.items()
    }
    metrics["per_document"] = {
        f"benchmark.{key}": dict(value) for key, value in metrics["per_format"].items()
    }
    answerable_correct = metrics["factual_answerable_cases"] - metrics["false_no_evidence"] - metrics["wrong"]
    metrics["supported_factual_recall"] = answerable_correct / metrics["factual_answerable_cases"]
    return metrics


if __name__ == "__main__":
    print(run_benchmark())
