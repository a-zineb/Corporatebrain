"""DOCX WordprocessingML reader that reconstructs a positional logical grid."""

from __future__ import annotations

from dataclasses import replace
from hashlib import sha256
from io import BytesIO
import re
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.oxml.ns import qn

from .models import LogicalCell, LogicalTable


def _clean(value: str) -> str:
    return re.sub(r"[ \t\r\f\v]+", " ", value.replace("\xa0", " ")).strip()


_SECRET_LINE = re.compile(r"(?im)^(\s*(?:password|passwd|token|api[_ -]?key|secret|private\s+key)\s*[:=]\s*).*$")


def _redact_paragraph(value: str) -> str:
    return _SECRET_LINE.sub(r"\1[REDACTED]", value)


def _paragraph_texts(tc) -> tuple[str, ...]:
    paragraphs: list[str] = []
    for paragraph in tc.findall(qn("w:p")):
        parts: list[str] = []
        for node in paragraph.iter():
            tag = node.tag.rsplit("}", 1)[-1]
            if tag == "t":
                parts.append(node.text or "")
            elif tag == "br":
                parts.append("\n")
            elif tag == "tab":
                parts.append(" ")
        text = _clean("".join(parts))
        if text:
            paragraphs.append(_redact_paragraph(text))
    return tuple(paragraphs)


def _style_metadata(tc) -> dict[str, object]:
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return {}
    shading = tc_pr.find(qn("w:shd"))
    borders = tc_pr.find(qn("w:tcBorders"))
    return {
        "background_color": shading.get(qn("w:fill")) if shading is not None else None,
        "has_border": borders is not None,
    }


def _shape(rows: list[list[str]]) -> str:
    if not rows:
        return "UNKNOWN_TABLE"
    normalized = [[_clean(value).casefold() for value in row] for row in rows]
    header = set(normalized[0])
    first_column = {row[0] for row in normalized if row}
    if {"version", "date"}.issubset(header) and header.intersection({"auteur", "author"}):
        return "VERSION_HISTORY"
    if header.intersection({"param", "parameter"}) and header.intersection({"value", "valeur"}):
        return "PARAMETER_TABLE"
    matrix_fields = {"protocol", "host", "hostname", "username", "password", "filedirectory", "directory", "cdr format"}
    if len(rows[0]) >= 3 and first_column.intersection(matrix_fields):
        return "MATRIX"
    if len(rows[0]) >= 3 and any(value for value in normalized[0][1:]):
        return "RECTANGULAR_RECORD_TABLE"
    if max((len(row) for row in rows), default=0) == 2:
        return "KEY_VALUE"
    return "UNKNOWN_TABLE"


def _logical_table(tbl, table_index: int, section_path: tuple[str, ...]) -> LogicalTable:
    grid = tbl.find(qn("w:tblGrid"))
    declared_width = len(grid.findall(qn("w:gridCol"))) if grid is not None else 0
    mutable_rows: list[list[LogicalCell | None]] = []
    vertical_origins: dict[int, tuple[int, int]] = {}
    origins: dict[tuple[int, int], LogicalCell] = {}

    for row_index, tr in enumerate(tbl.findall(qn("w:tr"))):
        slots: list[LogicalCell | None] = [None] * declared_width
        column = 0
        for tc_index, tc in enumerate(tr.findall(qn("w:tc"))):
            while column < len(slots) and slots[column] is not None:
                column += 1
            tc_pr = tc.find(qn("w:tcPr"))
            span_node = tc_pr.find(qn("w:gridSpan")) if tc_pr is not None else None
            column_span = int(span_node.get(qn("w:val"), "1")) if span_node is not None else 1
            while len(slots) < column + column_span:
                slots.append(None)
            merge = tc_pr.find(qn("w:vMerge")) if tc_pr is not None else None
            merge_value = merge.get(qn("w:val"), "continue") if merge is not None else None
            paragraphs = _paragraph_texts(tc)
            text = "\n".join(paragraphs)
            source_id = sha256(f"table:{table_index}:row:{row_index}:tc:{tc_index}".encode()).hexdigest()

            for offset in range(column_span):
                logical_column = column + offset
                if merge is not None and merge_value != "restart" and logical_column in vertical_origins:
                    origin_key = vertical_origins[logical_column]
                    origin = origins[origin_key]
                    origins[origin_key] = replace(origin, row_span=origin.row_span + 1)
                    cell = replace(
                        origin, row_index=row_index, column_index=logical_column,
                        column_span=1, is_merged_continuation=True,
                    )
                else:
                    cell = LogicalCell(
                        row_index=row_index, column_index=logical_column, row_span=1,
                        column_span=column_span if offset == 0 else 1, text=text,
                        paragraphs=paragraphs, style_metadata=_style_metadata(tc),
                        is_merged_continuation=offset > 0, source_xml_identity=source_id,
                    )
                    if offset == 0:
                        origins[(row_index, logical_column)] = cell
                    if merge is not None and merge_value == "restart":
                        vertical_origins[logical_column] = (row_index, logical_column)
                    elif merge is None:
                        vertical_origins.pop(logical_column, None)
                slots[logical_column] = cell
            column += column_span
        mutable_rows.append(slots)

    width = max([declared_width, *(len(row) for row in mutable_rows)], default=0)
    frozen_rows: list[tuple[LogicalCell, ...]] = []
    for row_index, row in enumerate(mutable_rows):
        padded = row + [None] * (width - len(row))
        final_row: list[LogicalCell] = []
        for column_index, cell in enumerate(padded):
            if cell is None:
                cell = LogicalCell(row_index, column_index, 1, 1, "", (), {}, False, "")
            elif not cell.is_merged_continuation:
                cell = origins.get((cell.row_index, cell.column_index), cell)
            final_row.append(cell)
        frozen_rows.append(tuple(final_row))

    # Apply record-aware redaction before the logical table can enter the
    # canonical document or any cache.  This covers separate label/value cells
    # as well as matrix password rows.
    protected_rows: list[tuple[LogicalCell, ...]] = []
    for row in frozen_rows:
        secret_row = bool(row and re.search(
            r"(?i)password|passwd|token|api[_ -]?key|secret|private\s+key", row[0].text
        ))
        if secret_row:
            row = tuple(
                replace(cell, text="[REDACTED]", paragraphs=("[REDACTED]",))
                if column_index > 0 and cell.text else cell
                for column_index, cell in enumerate(row)
            )
        protected_rows.append(row)
    frozen_rows = protected_rows

    values = [[cell.text for cell in row] for row in frozen_rows]
    shape = _shape(values)
    section = section_path[-1] if section_path else None
    return LogicalTable(
        table_id=f"table-{table_index}", section=section, section_path=section_path,
        rows=tuple(frozen_rows), logical_columns=width, source_order=table_index,
        shape=shape,
        metadata={
            "declared_grid_columns": declared_width,
            "merged_cell_count": sum(cell.is_merged_continuation for row in frozen_rows for cell in row),
        },
    )


def read_docx_structure(document: str | Path | bytes) -> tuple[LogicalTable, ...]:
    """Read all tables in body order with their inherited heading path."""
    doc = Document(str(document)) if isinstance(document, (str, Path)) else Document(BytesIO(document))
    paragraph_by_xml = {id(paragraph._p): paragraph for paragraph in doc.paragraphs}
    headings: dict[int, str] = {}
    tables: list[LogicalTable] = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = paragraph_by_xml.get(id(child))
            if paragraph is None or not _clean(paragraph.text):
                continue
            style = paragraph.style.name if paragraph.style else ""
            match = re.search(r"(?i)heading\s*(\d+)", style)
            if match:
                level = int(match.group(1))
                headings[level] = _clean(paragraph.text)
                headings = {key: value for key, value in headings.items() if key <= level}
        elif tag == "tbl":
            path = tuple(value for _, value in sorted(headings.items()))
            tables.append(_logical_table(child, len(tables), path))
    return tuple(tables)
