"""Phase-1 structured DOCX extraction and preview (no indexing)."""

from __future__ import annotations

from dataclasses import asdict, dataclass, field
from hashlib import sha256
from pathlib import Path
import re
from typing import Any, Iterable

from docx import Document
from document_engine import LogicalTable, read_docx_structure


@dataclass
class NormalizedBlock:
    text: str
    block_type: str
    source_file: str
    location: str
    section: str | None = None
    sheet_name: str | None = None
    table_index: int | None = None
    row_index: int | None = None
    metadata: dict[str, object] = field(default_factory=dict)

    @property
    def block_id(self) -> str:
        """Stable identity independent of ingestion time or insertion order."""

        return sha256(
            f"{self.source_file}\0{self.location}\0{self.block_type}\0{self.text}".encode("utf-8")
        ).hexdigest()

    def to_preview(self) -> dict[str, Any]:
        payload = asdict(self)
        payload["block_id"] = self.block_id
        return payload


@dataclass
class NormalizedChunk:
    text: str
    source_file: str
    location: str
    block_type: str
    chunk_ordinal: int
    source_block_indices: tuple[int, ...]
    content_sha256: str
    section: str | None = None
    table_index: int | None = None
    row_index: int | None = None
    metadata: dict[str, object] = field(default_factory=dict)

    def to_preview(self) -> dict[str, Any]:
        return {
            "text": self.text,
            "source_file": self.source_file,
            "location": self.location,
            "block_type": self.block_type,
            "chunk_ordinal": self.chunk_ordinal,
            "source_block_indices": list(self.source_block_indices),
            "content_sha256": self.content_sha256,
            "section": self.section,
            "table_index": self.table_index,
            "row_index": self.row_index,
            "metadata": dict(self.metadata),
        }


_SECRET_FIELD = re.compile(r"(password|passwd|token|api\s*key|secret|private\s*key)", re.I)
_SAFE_SECRET_PLACEHOLDERS = frozenset({
    "[redacted]", "***", "******", "*******", "to be defined", "n/a", "na", "not defined", "tbd",
})
_SECRET_VALUE = re.compile(
    r"(?i)\b(?:password|passwd|token|api[_ -]?key|secret|private\s+key)\s*[:=]\s*(?P<value>[^|;,\n]+)"
)


def _clean(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\xa0", " ")).strip()


def _safe_value(label: str, value: str) -> str:
    return "[REDACTED]" if _SECRET_FIELD.search(label) else value


def _is_safe_secret_placeholder(value: str) -> bool:
    """Return true only for an exact, normalized non-secret placeholder."""
    normalized = _clean(value).casefold()
    if normalized.strip("*"):
        normalized = re.sub(r"\*+$", "", normalized).strip()
    return normalized in _SAFE_SECRET_PLACEHOLDERS


def _contains_unredacted_secret_value(text: str) -> bool:
    """Detect credential-labelled values while allowing exact safe placeholders."""
    for match in _SECRET_VALUE.finditer(str(text or "")):
        if not _is_safe_secret_placeholder(match.group("value")):
            return True
    return False


def _cell_text(cell: Any) -> str:
    return _clean(" ".join(p.text for p in cell.paragraphs))


def _table_rows(table: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        values = [_cell_text(cell) for cell in row.cells]
        rows.append(values)
    return rows


def _key_value_blocks(rows: list[list[str]], source_file: str, location: str, section: str | None, table_index: int) -> list[NormalizedBlock]:
    blocks: list[NormalizedBlock] = []
    for row_index, row in enumerate(rows):
        if len(row) < 2:
            continue
        label = row[0]
        values = [value for value in row[1:] if value]
        if not label or not values:
            continue
        text = f"{label} = {' | '.join(_safe_value(label, value) for value in values)}"
        blocks.append(NormalizedBlock(
            text=text, block_type="table_row", source_file=source_file, location=location,
            section=section, table_index=table_index, row_index=row_index,
            metadata={"table_shape": "key_value", "normalization_strategy": "key_value"},
        ))
        # Preserve a group record and also expose safe Label: Value lines from
        # multiline cells as separately traceable child fields.
        for value in values:
            for line in value.splitlines():
                match = re.match(r"^\s*([^:=]{1,80})\s*[:=]\s*(.+?)\s*$", line)
                if not match:
                    continue
                child_label, child_value = _clean(match.group(1)), _clean(match.group(2))
                blocks.append(NormalizedBlock(
                    text=f"{label} | {child_label} = {_safe_value(child_label, child_value)}",
                    block_type="table_row", source_file=source_file, location=location,
                    section=section, table_index=table_index, row_index=row_index,
                    metadata={"table_shape": "key_value", "normalization_strategy": "key_value",
                              "parent_field": label, "subfield": child_label},
                ))
    return blocks


_VERTICAL_FIELD_LABELS = re.compile(
    r"(?i)^(?:enrichissement|enrichment|normalisation|normalization|correlation|"
    r"database\s+lookups?|duplicate\s+udr\s+check|filtrage|filtering|"
    r"selection|sélection)$"
)


def _vertical_field_value(label: str, value: str) -> bool:
    """Accept only clear field labels paired with a concise adjacent value."""
    clean = _clean(value)
    return bool(_VERTICAL_FIELD_LABELS.match(_clean(label))) and bool(clean) and len(clean) <= 120 and "\n" not in clean


def _matrix_blocks(rows: list[list[str]], source_file: str, location: str, section: str | None, table_index: int) -> list[NormalizedBlock]:
    if not rows or len(rows[0]) < 3:
        return []
    headers = rows[0]
    # A matrix has a row-label column and named logical columns.
    if not headers[0] or sum(bool(value) for value in headers[1:]) < 2:
        return []
    records: list[list[str]] = [(["System name = " + header] if headers[0].casefold() == "system name" and header else []) for header in headers[1:]]
    for row in rows[1:]:
        label = row[0] if row else ""
        if not label:
            continue
        for index, value in enumerate(row[1:]):
            if index >= len(records) or not value:
                continue
            records[index].append(f"{label} = {_safe_value(label, value)}")
    blocks: list[NormalizedBlock] = []
    for column_index, values in enumerate(records):
        if not values:
            continue
        blocks.append(NormalizedBlock(
            text=" | ".join(values), block_type="table_row", source_file=source_file,
            location=location, section=section, table_index=table_index,
            row_index=column_index, metadata={
                "logical_column_index": column_index,
                "table_shape": "matrix",
                "normalization_strategy": "column_records",
                "column_header": headers[column_index + 1],
            },
        ))
    return blocks


def _row_record_blocks(rows: list[list[str]], source_file: str, location: str, section: str | None, table_index: int) -> list[NormalizedBlock]:
    """Normalize ordinary rectangular tables with stable header/value pairs."""
    blocks: list[NormalizedBlock] = []
    if not rows:
        return blocks
    headers = list(rows[0])
    normalized_headers = [_clean(value).casefold() for value in headers]
    for row_index, row in enumerate(rows[1:], start=1):
        normalized_row = [_clean(value).casefold() for value in row]
        # Repeated header rows are ignored conservatively only when the
        # populated cells match the original header sequence.
        if normalized_row[:len(normalized_headers)] == normalized_headers and any(normalized_headers):
            continue
        pairs: list[str] = []
        for column_index, value in enumerate(row):
            value = _clean(value)
            if not value:
                continue
            label = headers[column_index] if column_index < len(headers) and headers[column_index] else f"Column {column_index + 1}"
            pairs.append(f"{label} = {_safe_value(label, value)}")
        if not pairs:
            continue
        blocks.append(NormalizedBlock(
            text=" | ".join(pairs), block_type="table_row", source_file=source_file,
            location=location, section=section, table_index=table_index, row_index=row_index,
            metadata={"table_shape": "row_records", "normalization_strategy": "row_records", "column_headers": headers},
        ))
    return blocks


def _table_blocks(table: Any, source_file: str, section: str | None, table_index: int,
                  logical_table: LogicalTable | None = None) -> list[NormalizedBlock]:
    # Logical rows preserve blank coordinates and propagate only proven vertical
    # merges.  ``row.cells`` remains a compatibility fallback for callers that
    # pass an isolated python-docx table.
    rows = logical_table.values() if logical_table is not None else _table_rows(table)
    location = f"Table {table_index}"
    common_metadata = {
        "section_path": list(logical_table.section_path) if logical_table else ([section] if section else []),
        "logical_row_count": len(rows),
        "logical_column_count": logical_table.logical_columns if logical_table else max((len(row) for row in rows), default=0),
        "merged_cell_count": logical_table.metadata.get("merged_cell_count", 0) if logical_table else 0,
        "logical_table_shape": logical_table.shape if logical_table else "",
    }
    # Two-level headers: the lower header row owns the destination columns;
    # the upper merged label is retained as parent-header metadata only.
    if len(rows) >= 3 and len(rows[0]) >= 3:
        top_values = [value for value in rows[0][1:] if value]
        second_headers = rows[1]
        if top_values and len(set(value.casefold() for value in top_values)) == 1 and sum(bool(v) for v in second_headers[1:]) >= 2:
            blocks = _row_record_blocks([second_headers, *rows[2:]], source_file, location, section, table_index)
            for block in blocks:
                block.metadata.update(common_metadata)
                block.metadata["parent_header"] = top_values[0]
                block.metadata["table_shape"] = "two_level_header_matrix"
            return blocks
    # Wide tables with a named first row are transposed configuration matrices.
    if rows and len(rows[0]) >= 3 and (
        rows[0][0].casefold() in {"system name", "system", "destination"}
        or any((row and row[0].casefold() in {"protocol", "host", "hostname", "username", "password", "filedirectory", "directory"}) for row in rows[1:])
    ):
        matrix = _matrix_blocks(rows, source_file, location, section, table_index)
        if matrix:
            for block in matrix:
                block.metadata.update(common_metadata)
            return matrix
    if rows and len(rows[0]) >= 3:
        records = _row_record_blocks(rows, source_file, location, section, table_index)
        if records:
            for block in records:
                block.metadata.update(common_metadata)
            return records
    key_values = _key_value_blocks(rows, source_file, location, section, table_index)
    if key_values:
        for block in key_values:
            block.metadata.update(common_metadata)
        return key_values
    raw = " | ".join(" | ".join(value for value in row if value) for row in rows if any(row))
    if not raw:
        return []
    return [NormalizedBlock(
        text=raw, block_type="table_row", source_file=source_file, location=location,
        section=section, table_index=table_index,
        metadata={"table_shape": "ambiguous", "normalization_strategy": "raw_fallback", **common_metadata},
    )]


def extract_docx_blocks(document: str | Path | bytes, source_file: str | None = None) -> list[NormalizedBlock]:
    """Extract DOCX paragraphs, headings, and structured tables in document order."""

    if isinstance(document, (str, Path)):
        path = Path(document)
        source_file = source_file or path.name
        doc = Document(str(path))
    else:
        from io import BytesIO
        if not source_file:
            raise ValueError("source_file is required for byte input")
        doc = Document(BytesIO(document))

    logical_tables = read_docx_structure(document)

    blocks: list[NormalizedBlock] = []
    section: str | None = None
    table_index = 0
    body = doc.element.body
    children = list(body.iterchildren())
    skip_vertical_ids: set[int] = set()
    for child_index, child in enumerate(children):
        if id(child) in skip_vertical_ids:
            continue
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = next((p for p in doc.paragraphs if p._p is child), None)
            if paragraph is None or not _clean(paragraph.text):
                continue
            style = paragraph.style.name if paragraph.style else ""
            is_heading = bool(re.search(r"(?i)heading\s*\d+", style))
            if is_heading:
                section = _clean(paragraph.text)
            paragraph_text = _clean(paragraph.text)
            if child_index + 1 < len(children):
                next_child = children[child_index + 1]
                if next_child.tag.rsplit("}", 1)[-1] == "p":
                    next_paragraph = next((p for p in doc.paragraphs if p._p is next_child), None)
                    next_text = _clean(next_paragraph.text) if next_paragraph is not None else ""
                    if _vertical_field_value(paragraph_text, next_text):
                        if is_heading:
                            blocks.append(NormalizedBlock(
                                text=paragraph_text, block_type="heading", source_file=source_file,
                                location="Corps du document", section=section,
                            ))
                        paragraph_text = f"{paragraph_text} = {_safe_value(paragraph_text, next_text)}"
                        skip_vertical_ids.add(id(next_child))
                        blocks.append(NormalizedBlock(
                            text=paragraph_text, block_type="table_row", source_file=source_file,
                            location="Corps du document", section=section, row_index=child_index,
                            metadata={"table_shape": "vertical_key_value", "normalization_strategy": "key_value", "value_paragraph_index": child_index + 1},
                        ))
                        continue
            blocks.append(NormalizedBlock(
                text=paragraph_text, block_type="heading" if is_heading else "paragraph",
                source_file=source_file, location="Corps du document", section=section,
            ))
        elif tag == "tbl":
            table = next((t for t in doc.tables if t._tbl is child), None)
            if table is not None:
                logical = logical_tables[table_index] if table_index < len(logical_tables) else None
                blocks.extend(_table_blocks(table, source_file, section, table_index, logical))
                table_index += 1
    return blocks


def preview_docx(document: str | Path | bytes, source_file: str | None = None) -> list[dict[str, Any]]:
    """Return JSON-friendly normalized blocks without indexing or persistence."""

    return [block.to_preview() for block in extract_docx_blocks(document, source_file)]


def _sentences(text: str) -> list[str]:
    return [part.strip() for part in re.split(r"(?<=[.!?])\s+", text) if part.strip()]


def normalized_blocks_to_chunks(
    blocks: Iterable[NormalizedBlock], *, max_length: int = 1000, overlap: int = 250
) -> list[NormalizedChunk]:
    """Convert normalized blocks into deterministic, structure-aware chunks.

    Structured table records are always emitted as one atomic chunk. Narrative
    paragraphs are split by sentence and may overlap at chunk boundaries.
    """

    if max_length <= 0 or overlap < 0:
        raise ValueError("max_length must be positive and overlap must be non-negative")
    chunks: list[NormalizedChunk] = []

    def emit(text: str, block: NormalizedBlock, index: int, source_indices: tuple[int, ...], block_type: str | None = None) -> None:
        clean = text.strip()
        if not clean:
            return
        chunks.append(NormalizedChunk(
            text=clean,
            source_file=block.source_file,
            location=block.location,
            block_type=block_type or block.block_type,
            chunk_ordinal=index,
            source_block_indices=source_indices,
            content_sha256=sha256(clean.encode("utf-8")).hexdigest(),
            section=block.section,
            table_index=block.table_index,
            row_index=block.row_index,
            metadata=dict(block.metadata),
        ))

    for block_index, block in enumerate(blocks):
        if block.block_type == "table_row" or block.metadata.get("normalization_strategy") in {
            "key_value", "column_records", "row_records", "raw_fallback"
        }:
            emit(block.text, block, len(chunks), (block_index,))
            continue
        if block.block_type == "heading":
            emit(block.text, block, len(chunks), (block_index,))
            continue
        sentences = _sentences(block.text)
        if not sentences:
            continue
        current = ""
        for sentence in sentences:
            if current and len(current) + 1 + len(sentence) > max_length:
                emit(current, block, len(chunks), (block_index,))
                tail = current[-overlap:] if overlap else ""
                current = (tail + " " + sentence).strip() if tail else sentence
            else:
                current = f"{current} {sentence}".strip()
        emit(current, block, len(chunks), (block_index,))
    return chunks


def preview_docx_chunks(document: str | Path | bytes, source_file: str | None = None, *, max_length: int = 1000, overlap: int = 250) -> list[dict[str, Any]]:
    """Preview normalized DOCX chunks without indexing or persistence."""

    blocks = extract_docx_blocks(document, source_file)
    return [chunk.to_preview() for chunk in normalized_blocks_to_chunks(blocks, max_length=max_length, overlap=overlap)]


def build_structured_docx_index_payload(
    file_bytes: bytes,
    source_file: str,
    *,
    file_hash: str,
    geographical_entity: str = "",
    application: str = "",
) -> dict[str, list[Any]]:
    """Prepare a complete structured DOCX index payload without writing it."""

    chunks = preview_docx_chunks(file_bytes, source_file)
    ids: list[str] = []
    documents: list[str] = []
    metadatas: list[dict[str, Any]] = []
    for chunk in chunks:
        chunk_id = f"{file_hash}_chunk_{chunk['chunk_ordinal']}"
        if chunk_id in ids:
            raise ValueError("duplicate deterministic structured chunk ID")
        text = chunk["text"]
        if _contains_unredacted_secret_value(text):
            raise ValueError("unredacted secret-like value in structured payload")
        ids.append(chunk_id)
        documents.append(text)
        metadata = {
            "source_file": source_file,
            "saved_as": source_file,
            "file_hash": file_hash,
            "geographical_entity": geographical_entity,
            "application": application,
            "block_type": chunk["block_type"],
            "section": chunk.get("section") or "",
            "location": chunk["location"],
            "chunk_ordinal": chunk["chunk_ordinal"],
            "content_sha256": chunk["content_sha256"],
        }
        for key in ("table_index", "row_index"):
            if chunk.get(key) is not None:
                metadata[key] = chunk[key]
        metadatas.append(metadata)
    return {"ids": ids, "documents": documents, "metadatas": metadatas}
