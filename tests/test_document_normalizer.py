import hashlib
import io
import zipfile

import fitz
from docx import Document
from openpyxl import Workbook

from document_normalizer import normalize_document


def _docx_bytes():
    stream = io.BytesIO()
    document = Document()
    document.add_heading("Collection", level=1)
    document.add_paragraph("MediationZone performs processing.")
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text, table.cell(0, 1).text = "Host", "10.0.0.1"
    table.cell(1, 0).text, table.cell(1, 1).text = "Password", "super-secret"
    table.cell(2, 0).text, table.cell(2, 1).text = "Frequency", "5 minutes"
    document.save(stream)
    return stream.getvalue()


def test_docx_is_deterministic_structured_and_redacted():
    data = _docx_bytes()
    first = normalize_document(data, "sample.docx")
    second = normalize_document(data, "sample.docx")
    assert first == second
    assert first.file_hash == hashlib.sha256(data).hexdigest()
    assert [block.block_id for block in first.blocks] == [block.block_id for block in second.blocks]
    assert any(block.section == "Collection" for block in first.blocks)
    text = first.canonical_text()
    assert "Host = 10.0.0.1" in text
    assert "Password = [REDACTED]" in text
    assert "super-secret" not in text


def test_pdf_preserves_page_and_does_not_silently_accept_empty():
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Protocol = SFTP")
    data = pdf.tobytes()
    pdf.close()
    result = normalize_document(data, "sample.pdf")
    assert result.blocks[0].page == 1
    assert "Protocol = SFTP" in result.canonical_text()

    empty = fitz.open()
    empty.new_page()
    empty_data = empty.tobytes()
    empty.close()
    assert "OCR may be required" in normalize_document(empty_data, "empty.pdf").warnings[0]


def test_xlsx_preserves_sheet_rows_columns_and_empty_position():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Systems"
    sheet.append(["System name", "Protocol", "Host", "FileDirectory"])
    sheet.append(["BI", "SFTP", "172.26.60.12", None])
    stream = io.BytesIO()
    workbook.save(stream)
    result = normalize_document(stream.getvalue(), "systems.xlsx")
    block = result.blocks[0]
    assert block.sheet == "Systems" and block.row_index == 2
    assert block.metadata["cell_range"] == "A2:D2"
    assert block.metadata["sheet_visibility"] == "visible"
    assert block.metadata["column_headers"][3] == "FileDirectory"
    assert "System name = BI | Protocol = SFTP | Host = 172.26.60.12" == block.text


def test_xlsx_matrix_and_key_value_relationships_are_preserved():
    workbook = Workbook()
    matrix = workbook.active
    matrix.title = "Matrix"
    matrix.append(["System name", "BI", "DWH"])
    matrix.append(["Protocol", "SFTP", "FTP"])
    matrix.append(["Host", "10.0.0.1", "10.0.0.2"])
    key_values = workbook.create_sheet("Config")
    key_values.append(["Collection Frequency", "5 minutes"])
    key_values.append(["Password", "do-not-leak"])
    stream = io.BytesIO()
    workbook.save(stream)
    result = normalize_document(stream.getvalue(), "matrix.xlsx")
    assert any(block.text == "System name = BI | Protocol = SFTP | Host = 10.0.0.1" for block in result.blocks)
    assert any(block.text == "Collection Frequency = 5 minutes" and block.block_type == "key_value" for block in result.blocks)
    assert "do-not-leak" not in result.canonical_text()


def test_xlsx_hidden_sheet_is_indexed_with_visibility_and_coordinates():
    workbook = Workbook()
    workbook.active.title = "Visible"
    hidden = workbook.create_sheet("STAT")
    hidden.sheet_state = "hidden"
    hidden.append(["Metric", "Value"])
    hidden.append(["Tests", 7])
    stream = io.BytesIO()
    workbook.save(stream)
    result = normalize_document(stream.getvalue(), "testbook.xlsx")
    stat = next(block for block in result.blocks if block.sheet == "STAT" and "Tests = 7" in block.text)
    assert stat.metadata["sheet_visibility"] == "hidden"
    assert stat.metadata["cell_range"] == "A2:B2"


def test_xlsx_preserves_merged_values_formulas_and_multiline_content():
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Group", "Value", "Formula"])
    sheet.append(["Merged group", "line one\nline two", "=1+1"])
    sheet.append([None, "next", None])
    sheet.merge_cells("A2:A3")
    stream = io.BytesIO()
    workbook.save(stream)
    result = normalize_document(stream.getvalue(), "structured.xlsx")
    second = next(block for block in result.blocks if block.row_index == 3)
    assert "Group = Merged group" in second.text
    first = next(block for block in result.blocks if block.row_index == 2)
    assert "line one line two" in first.text
    assert first.metadata["formulas"]["C2"] == "=1+1"
    assert "A2:A3" in first.metadata["merged_ranges"]


def test_csv_headers_rows_identity_and_no_cross_document_mixing():
    a = normalize_document(b"System,Host,Password\nBI,1.1.1.1,raw-secret\n", "a.csv")
    b = normalize_document(b"System,Host\nBI,2.2.2.2\n", "b.csv")
    assert a.file_hash != b.file_hash
    assert all(block.file_hash == a.file_hash and block.source_file == "a.csv" for block in a.blocks)
    assert "Host = 1.1.1.1" in a.blocks[0].text
    assert "Password = [REDACTED]" in a.blocks[0].text and "raw-secret" not in a.canonical_text()
    assert a.blocks[0].metadata["row_start"] == 2
    assert a.blocks[0].metadata["row_end"] == 2
    assert all(block.block_id not in {other.block_id for other in b.blocks} for block in a.blocks)


def test_doc_conversion_uses_original_bytes_for_identity():
    original = b"legacy binary bytes"
    converted = _docx_bytes()
    result = normalize_document(original, "legacy.doc", doc_converter=lambda *_: converted)
    assert result.file_hash == hashlib.sha256(original).hexdigest()
    assert result.blocks and all(block.file_hash == result.file_hash for block in result.blocks)


def test_zip_dispatch_is_safe_and_uses_outer_identity():
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("nested/data.csv", "Name,Port\nAPI,443\n")
        archive.writestr("../escape.csv", "Bad,Value\nX,Y\n")
        archive.writestr("script.exe", b"danger")
    result = normalize_document(stream.getvalue(), "bundle.zip")
    assert len(result.blocks) == 1
    assert result.blocks[0].source_file == "bundle.zip::nested/data.csv"
    assert result.blocks[0].file_hash == result.file_hash
    assert "Port = 443" in result.blocks[0].text
    assert len(result.warnings) == 2


def test_repeated_normalization_is_identical_across_non_docx_formats():
    pdf = fitz.open()
    pdf.new_page().insert_text((72, 72), "Version = 2.0")
    pdf_bytes = pdf.tobytes()
    pdf.close()

    workbook = Workbook()
    workbook.active.append(["Name", "Value", "Status"])
    workbook.active.append(["Mode", "PUSH", "Active"])
    xlsx_stream = io.BytesIO()
    workbook.save(xlsx_stream)

    fixtures = [
        (pdf_bytes, "repeat.pdf"),
        (xlsx_stream.getvalue(), "repeat.xlsx"),
        (b"Name,Value\nMode,PUSH\n", "repeat.csv"),
    ]
    for data, name in fixtures:
        first = normalize_document(data, name)
        second = normalize_document(data, name)
        assert first == second
        assert first.canonical_text() == second.canonical_text()
