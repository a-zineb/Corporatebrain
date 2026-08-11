from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from document_engine import read_docx_structure
from document_normalizer import normalize_document


def _bytes(document: Document) -> bytes:
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_horizontal_and_vertical_merges_preserve_grid_and_parent_context():
    document = Document()
    table = document.add_table(rows=4, cols=4)
    for column, value in enumerate(("Flux", "PARAM", "VALUE", "DESCRIPTION")):
        table.cell(0, column).text = value
    table.cell(1, 0).text = "Tango"
    table.cell(1, 0).merge(table.cell(3, 0))
    for row, destination in enumerate(("SVRCRA", "DWH", "BigData"), start=1):
        table.cell(row, 1).text = f"TANGO_TO_{destination}"
        table.cell(row, 2).text = "Y/N"
    logical = read_docx_structure(_bytes(document))[0]
    assert logical.logical_columns == 4
    assert [row[0] for row in logical.values()] == ["Flux", "Tango", "Tango", "Tango"]
    assert logical.rows[2][0].is_merged_continuation
    assert logical.metadata["merged_cell_count"] == 2


def test_blank_coordinates_are_not_shifted():
    document = Document()
    table = document.add_table(rows=2, cols=4)
    for column, value in enumerate(("A", "B", "C", "D")):
        table.cell(0, column).text = value
    table.cell(1, 0).text = "x"
    table.cell(1, 2).text = "y"
    table.cell(1, 3).text = "z"
    assert read_docx_structure(_bytes(document))[0].values()[1] == ["x", "", "y", "z"]


def test_multiline_cell_creates_safe_subfields_and_never_keeps_secret():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Server Details"
    table.cell(0, 1).text = "IP : 10.0.0.1\nLogin : user\nPassword : real-secret"
    canonical = normalize_document(_bytes(document), "server.docx")
    text = canonical.canonical_text()
    assert "IP = 10.0.0.1" in text
    assert "Login = user" in text
    assert "real-secret" not in text
    assert "Password : [REDACTED]" in text


def test_style_changes_do_not_change_logical_values():
    base = Document()
    table = base.add_table(rows=2, cols=3)
    for column, value in enumerate(("System", "DWH", "BIG DATA")):
        table.cell(0, column).text = value
    for column, value in enumerate(("Host", "10.0.0.1", "10.0.0.2")):
        table.cell(1, column).text = value
    changed = Document(BytesIO(_bytes(base)))
    for row in changed.tables[0].rows:
        for cell in row.cells:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "00FF00")
            cell._tc.get_or_add_tcPr().append(shading)
            for run in cell.paragraphs[0].runs:
                run.bold = not bool(run.bold)
    first = read_docx_structure(_bytes(base))[0].values()
    second = read_docx_structure(_bytes(changed))[0].values()
    assert first == second


def test_real_tango_visual_grid_assertions():
    path = next(Path("doc_storage_v2").glob("*Tango*.docx"))
    tables = read_docx_structure(path)
    assert [(len(table.rows), table.logical_columns) for table in tables] == [
        (3, 2), (3, 4), (28, 2), (3, 3), (3, 4), (9, 2), (8, 4), (4, 4)
    ]
    distribution = tables[6].values()
    assert distribution[0] == ["System name", "DWH", "BIG DATA", "FTP_CRA"]
    assert distribution[2] == ["Host", "172.21.75.61", "172.26.60.12", "172.21.75.61"]
    parameters = tables[7].values()
    assert [row[0] for row in parameters[1:]] == ["Tango", "Tango", "Tango"]


def test_docx_xlsx_csv_pdf_matrix_parity():
    import fitz
    from openpyxl import Workbook

    rows = [["System name", "DWH", "BIG DATA"], ["Host", "10.0.0.1", "10.0.0.2"],
            ["Protocol", "SFTP", "SFTP"]]

    word = Document()
    table = word.add_table(rows=3, cols=3)
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value

    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(row)
    xlsx_stream = BytesIO()
    workbook.save(xlsx_stream)

    pdf = fitz.open()
    page = pdf.new_page(width=360, height=180)
    left, top, width, height = 20, 20, 320, 120
    for index in range(4):
        x = left + index * width / 3
        y = top + index * height / 3
        page.draw_line((x, top), (x, top + height))
        page.draw_line((left, y), (left + width, y))
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            page.insert_text((left + column_index * width / 3 + 4,
                              top + row_index * height / 3 + 20), value, fontsize=8)
    pdf_bytes = pdf.tobytes()
    pdf.close()

    documents = [
        normalize_document(_bytes(word), "matrix.docx"),
        normalize_document(xlsx_stream.getvalue(), "matrix.xlsx"),
        normalize_document("\n".join(",".join(row) for row in rows).encode(), "matrix.csv"),
        normalize_document(pdf_bytes, "matrix.pdf"),
    ]
    expected = {
        "System name = DWH | Host = 10.0.0.1 | Protocol = SFTP",
        "System name = BIG DATA | Host = 10.0.0.2 | Protocol = SFTP",
    }
    for document in documents:
        records = {block.text for block in document.blocks if block.block_type == "table_row"}
        assert expected <= records


def test_borderless_pdf_uses_coordinate_table_fallback():
    import fitz

    pdf = fitz.open()
    page = pdf.new_page(width=360, height=220)
    rows = [("System name", "DWH", "BI"), ("Host", "10.0.0.1", "10.0.0.2"),
            ("Protocol", "SFTP", "SFTP")]
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            page.insert_text((30 + column_index * 110, 50 + row_index * 30), value, fontsize=10)
    structure = __import__("document_engine").read_pdf_structure(pdf.tobytes())
    pdf.close()
    assert len(structure.logical_tables) == 1
    table = structure.logical_tables[0]
    assert table.metadata["coordinate_reconstruction"] is True
    assert table.values() == [list(row) for row in rows]
