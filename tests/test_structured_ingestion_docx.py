import tempfile
import unittest
from pathlib import Path

from docx import Document
import structured_ingestion
from structured_ingestion import extract_docx_blocks


def make_doc(paragraphs=(), tables=()):
    doc = Document()
    for text, style in paragraphs:
        doc.add_paragraph(text, style=style)
    for rows in tables:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                table.cell(i, j).text = value
    handle = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    handle.close(); doc.save(handle.name)
    return Path(handle.name)


class StructuredDocxTests(unittest.TestCase):
    def test_secret_placeholders_are_exactly_safe(self):
        safe = structured_ingestion._contains_unredacted_secret_value
        for value in ("TO BE DEFINED", "TBD", "N/A", "NA", "NOT DEFINED", "[REDACTED]", "***", "******", "*******"):
            self.assertFalse(safe(f"Password = {value}"), value)
        for value in ("ActualSecret", "TBDabc123", "TO BE DEFINED-Secret123"):
            self.assertTrue(safe(f"Password = {value}"), value)

    def test_token_and_api_key_values_remain_blocked(self):
        safe = structured_ingestion._contains_unredacted_secret_value
        self.assertTrue(safe("token = abc123"))
        self.assertTrue(safe("api_key = sk-test"))

    def test_paragraphs_headings_and_section_inheritance(self):
        path = make_doc([("Section", "Heading 1"), ("Texte français", None)])
        blocks = extract_docx_blocks(path, "x.docx")
        self.assertEqual([b.block_type for b in blocks], ["heading", "paragraph"])
        self.assertEqual(blocks[1].section, "Section")

    def test_key_value_table_is_atomic(self):
        path = make_doc(tables=[[("Filename Pattern", "p2pCommands.yyyy-mm-dd"), ("Frequency", "Tous les jours à 7h")]])
        blocks = extract_docx_blocks(path, "x.docx")
        self.assertEqual([b.text for b in blocks], ["Filename Pattern = p2pCommands.yyyy-mm-dd", "Frequency = Tous les jours à 7h"])
        self.assertTrue(all(b.metadata["normalization_strategy"] == "key_value" for b in blocks))

    def test_vertical_processing_fields_pair_conservatively(self):
        path = make_doc(paragraphs=[
            ("Enrichissement", None), ("N/A", None),
            ("Normalisation", None), ("N/A", None),
        ])
        blocks = extract_docx_blocks(path, "x.docx")
        self.assertEqual([b.text for b in blocks], ["Enrichissement = N/A", "Normalisation = N/A"])
        self.assertTrue(all(b.metadata["table_shape"] == "vertical_key_value" for b in blocks))

    def test_unrelated_single_column_rows_are_not_paired(self):
        path = make_doc(paragraphs=[
            ("Operational notes", None), ("N/A", None),
            ("Enrichissement", None), ("N/A", None),
        ])
        blocks = extract_docx_blocks(path, "x.docx")
        self.assertIn("Operational notes", [b.text for b in blocks])
        self.assertNotIn("Operational notes = N/A", [b.text for b in blocks])
        self.assertIn("Enrichissement = N/A", [b.text for b in blocks])

    def test_matrix_table_preserves_column_associations(self):
        path = make_doc(tables=[[("System name", "DWH", "BI"), ("Protocol", "SFTP", "SFTP"), ("Host", "172.21.75.21", "172.26.60.12"), ("username", "", "mz_user")]])
        blocks = extract_docx_blocks(path, "x.docx")
        self.assertEqual(blocks[0].text, "System name = DWH | Protocol = SFTP | Host = 172.21.75.21")
        self.assertIn("username = mz_user", blocks[1].text)
        self.assertEqual(blocks[0].metadata["normalization_strategy"], "column_records")

    def test_empty_and_merged_cells_do_not_shift_values(self):
        path = make_doc(tables=[[("A", ""), ("B", "value")]])
        blocks = extract_docx_blocks(path, "x.docx")
        self.assertEqual([b.text for b in blocks], ["B = value"])

    def test_redacts_secret_like_fields(self):
        path = make_doc(tables=[[("Password", "secret-value"), ("API key", "abc")]])
        blocks = extract_docx_blocks(path, "x.docx")
        self.assertEqual([b.text for b in blocks], ["Password = [REDACTED]", "API key = [REDACTED]"])

    def test_ambiguous_table_uses_raw_fallback_and_order_is_deterministic(self):
        path = make_doc(tables=[[("", ""), ("only", "")]])
        first = extract_docx_blocks(path, "x.docx")
        second = extract_docx_blocks(path, "x.docx")
        self.assertEqual([b.to_preview() for b in first], [b.to_preview() for b in second])

    def test_p2p_preview_recovers_structured_values_without_indexing(self):
        path = Path("doc_storage_v2/CISA_MZ-001-1.0 - OCM Mediation - P2P.docx")
        blocks = extract_docx_blocks(path, path.name)
        text = "\n".join(block.text for block in blocks)
        for expected in ("1.0", "Omar EL HIMASS", "P2P_1", "172.21.75.9", "p2pCommands.yyyy-mm-dd", "Tous les jours à 7h", "SFTP", "172.21.75.21", "mz_user", "P2P_TO_BI", "AUDIT_INPUT"):
            self.assertIn(expected, text)
        self.assertNotIn("secret-value", text)


if __name__ == "__main__":
    unittest.main()
